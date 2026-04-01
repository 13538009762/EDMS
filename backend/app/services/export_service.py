"""Export TipTap/ProseMirror JSON to DOCX and PDF."""
from __future__ import annotations

import json
import re
from html import escape as html_escape
from io import BytesIO
from typing import Any

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from xhtml2pdf import pisa


def _walk_text(node: dict[str, Any], parts: list[str]) -> None:
    ntype = node.get("type")
    if ntype == "text":
        t = node.get("text") or ""
        marks = node.get("marks") or []
        if any(m.get("type") == "bold" for m in marks):
            t = f"<b>{html_escape(t)}</b>"
        else:
            t = html_escape(t)
        if any(m.get("type") == "italic" for m in marks):
            t = f"<i>{t}</i>"
        if any(m.get("type") == "underline" for m in marks):
            t = f"<u>{t}</u>"
        parts.append(t)
        return
    for ch in node.get("content") or []:
        _walk_text(ch, parts)


def tiptap_json_to_plain(doc_json: str) -> str:
    try:
        root = json.loads(doc_json)
    except json.JSONDecodeError:
        return ""
    parts: list[str] = []

    def block(node: dict[str, Any]) -> None:
        t = node.get("type")
        if t == "paragraph":
            line: list[str] = []
            for ch in node.get("content") or []:
                _walk_text(ch, line)
            parts.append("".join(line))
        elif t in ("heading",):
            line = []
            for ch in node.get("content") or []:
                _walk_text(ch, line)
            lvl = (node.get("attrs") or {}).get("level") or 1
            parts.append(f"{'#' * int(lvl)} " + "".join(line))
        elif t == "bulletList" or t == "orderedList":
            for item in node.get("content") or []:
                if item.get("type") == "listItem":
                    sub: list[str] = []
                    for ch in item.get("content") or []:
                        if ch.get("type") == "paragraph":
                            for c2 in ch.get("content") or []:
                                _walk_text(c2, sub)
                    parts.append("- " + "".join(sub))
        elif t == "doc":
            for ch in node.get("content") or []:
                block(ch)
        else:
            for ch in node.get("content") or []:
                block(ch)

    block(root)
    return "\n".join(parts)


def tiptap_json_to_html(doc_json: str) -> str:
    try:
        root = json.loads(doc_json)
    except json.JSONDecodeError:
        return "<p></p>"

    def inline(node: dict[str, Any]) -> str:
        if node.get("type") == "text":
            t = html_escape(node.get("text") or "")
            for m in node.get("marks") or []:
                mt = m.get("type")
                if mt == "bold":
                    t = f"<strong>{t}</strong>"
                elif mt == "italic":
                    t = f"<em>{t}</em>"
                elif mt == "underline":
                    t = f"<u>{t}</u>"
                elif mt == "strike":
                    t = f"<s>{t}</s>"
                elif mt == "textStyle":
                    attrs = m.get("attrs") or {}
                    color = attrs.get("color")
                    if color:
                        t = f'<span style="color:{html_escape(color)}">{t}</span>'
            return t
        return ""

    def block(node: dict[str, Any]) -> str:
        t = node.get("type")
        if t == "paragraph":
            inner = "".join(
                inline(c) if c.get("type") == "text" else "" for c in node.get("content") or []
            )
            style = ""
            ta = (node.get("attrs") or {}).get("textAlign")
            if ta:
                style = f' style="text-align:{html_escape(ta)}"'
            return f"<p{style}>{inner or '&nbsp;'}</p>"
        if t == "heading":
            lvl = int((node.get("attrs") or {}).get("level") or 1)
            inner = "".join(
                inline(c) if c.get("type") == "text" else "" for c in node.get("content") or []
            )
            return f"<h{lvl}>{inner}</h{lvl}>"
        if t == "bulletList":
            items = "".join(block(c) for c in node.get("content") or [])
            return f"<ul>{items}</ul>"
        if t == "orderedList":
            items = "".join(block(c) for c in node.get("content") or [])
            return f"<ol>{items}</ol>"
        if t == "listItem":
            inner = "".join(block(c) for c in node.get("content") or [])
            return f"<li>{inner}</li>"
        if t == "table":
            inner = "".join(block(c) for c in node.get("content") or [])
            return f"<table>{inner}</table>"
        if t == "tableRow":
            inner = "".join(block(c) for c in node.get("content") or [])
            return f"<tr>{inner}</tr>"
        if t in ("tableCell", "tableHeader"):
            tag = "th" if t == "tableHeader" else "td"
            inner = "".join(block(c) for c in node.get("content") or [])
            colspan = (node.get("attrs") or {}).get("colspan", 1)
            rowspan = (node.get("attrs") or {}).get("rowspan", 1)
            attrs_str = f' colspan="{colspan}"' if colspan > 1 else ""
            attrs_str += f' rowspan="{rowspan}"' if rowspan > 1 else ""
            return f"<{tag}{attrs_str}>{inner}</{tag}>"
        if t == "image":
            src = (node.get("attrs") or {}).get("src") or ""
            return f"<img src='{html_escape(src)}' />"
        if t == "doc":
            return "".join(block(c) for c in node.get("content") or [])
        return "".join(block(c) for c in node.get("content") or [])

    body = block(root)
    # xhtml2pdf requires a font reference to render CJK characters
    head = """
    <head>
    <meta charset='utf-8'>
    <style>
      @font-face {
        font-family: 'msyh';
        src: url('c:/windows/fonts/msyh.ttc');
      }
      body {
        font-family: 'msyh', sans-serif;
        font-size: 14px;
        line-height: 1.5;
        color: #333;
      }
      table {
        width: 100%;
        border-collapse: collapse;
        margin-top: 10px;
        margin-bottom: 10px;
      }
      th, td {
        border: 1px solid #cccccc;
        padding: 5px;
      }
      th {
        background-color: #f5f5f5;
      }
      img {
        max-width: 100%;
      }
    </style>
    </head>
    """
    return f"<!DOCTYPE html><html>{head}<body>{body}</body></html>"


def export_docx_bytes(doc_json: str) -> bytes:
    """Build a simple DOCX from TipTap JSON."""
    data = json.loads(doc_json) if doc_json else {"type": "doc", "content": []}
    d = DocxDocument()

    def add_paragraph_from_node(node: dict[str, Any]) -> None:
        t = node.get("type")
        if t == "paragraph":
            p = d.add_paragraph()
            for ch in node.get("content") or []:
                if ch.get("type") != "text":
                    continue
                run = p.add_run(ch.get("text") or "")
                for m in ch.get("marks") or []:
                    mt = m.get("type")
                    if mt == "bold":
                        run.bold = True
                    elif mt == "italic":
                        run.italic = True
                    elif mt == "underline":
                        run.underline = True
                    elif mt == "strike":
                        run.font.strike = True
                    elif mt == "textStyle":
                        col = (m.get("attrs") or {}).get("color")
                        if col and col.startswith("#") and len(col) >= 7:
                            r = int(col[1:3], 16)
                            g = int(col[3:5], 16)
                            b = int(col[5:7], 16)
                            run.font.color.rgb = RGBColor(r, g, b)
            ta = (node.get("attrs") or {}).get("textAlign")
            if ta == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif ta == "right":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif t == "heading":
            lvl = int((node.get("attrs") or {}).get("level") or 1)
            p = d.add_heading(level=min(lvl, 3))
            for ch in node.get("content") or []:
                if ch.get("type") == "text":
                    p.add_run(ch.get("text") or "")
        elif t in ("bulletList", "orderedList"):
            style = "List Bullet" if t == "bulletList" else "List Number"
            for item in node.get("content") or []:
                if item.get("type") != "listItem":
                    continue
                for sub in item.get("content") or []:
                    if sub.get("type") == "paragraph":
                        tp = d.add_paragraph(style=style)
                        for ch in sub.get("content") or []:
                            if ch.get("type") == "text":
                                tp.add_run(ch.get("text") or "")
        elif t == "doc":
            for ch in node.get("content") or []:
                add_paragraph_from_node(ch)

    add_paragraph_from_node(data)
    buf = BytesIO()
    d.save(buf)
    return buf.getvalue()


def export_pdf_bytes(doc_json: str) -> bytes:
    html = tiptap_json_to_html(doc_json or "{}")
    out = BytesIO()
    pisa.CreatePDF(src=html, dest=out, encoding="utf-8")
    return out.getvalue()


def apply_punctuation_fixes(text: str) -> str:
    """Lightweight Chinese punctuation normalization."""
    text = re.sub(r"，{2,}", "，", text)
    text = re.sub(r"。{2,}", "。", text)
    text = text.replace("。。", "。")
    return text
